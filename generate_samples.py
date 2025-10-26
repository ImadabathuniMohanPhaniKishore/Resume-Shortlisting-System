from docx import Document
import os

def create_sample_resumes():
    resumes_data = [
        {
            "name": "John Smith",
            "filename": "john_smith_resume",
            "email": "john.smith@email.com",
            "phone": "(555) 123-4567",
            "experience": "5 years of experience",
            "skills": ["Python", "Flask", "React", "PostgreSQL", "Docker", "AWS", "Machine Learning", "REST API"],
            "education": "B.S. Computer Science, MIT, 2018",
            "summary": "Full Stack Developer with expertise in building scalable web applications"
        },
        {
            "name": "Sarah Johnson",
            "filename": "sarah_johnson_resume",
            "email": "sarah.j@email.com",
            "phone": "(555) 234-5678",
            "experience": "3 years of experience",
            "skills": ["JavaScript", "Node.js", "React", "Angular", "MongoDB", "Express", "Git", "Agile"],
            "education": "B.S. Software Engineering, Stanford, 2020",
            "summary": "Frontend developer specializing in modern JavaScript frameworks"
        },
        {
            "name": "Michael Chen",
            "filename": "michael_chen_resume",
            "email": "m.chen@email.com",
            "phone": "(555) 345-6789",
            "experience": "7 years of experience",
            "skills": ["Python", "Django", "Data Science", "TensorFlow", "NLP", "SQL", "Tableau", "Scikit-learn"],
            "education": "M.S. Data Science, Berkeley, 2016",
            "summary": "Data Scientist with strong background in machine learning and NLP"
        },
        {
            "name": "Emily Rodriguez",
            "filename": "emily_rodriguez_resume",
            "email": "emily.r@email.com",
            "phone": "(555) 456-7890",
            "experience": "2 years of experience",
            "skills": ["HTML", "CSS", "JavaScript", "Bootstrap", "jQuery", "Photoshop", "Figma"],
            "education": "B.A. Graphic Design, UCLA, 2021",
            "summary": "UI/UX Designer with frontend development skills"
        },
        {
            "name": "David Kim",
            "filename": "david_kim_resume",
            "email": "david.kim@email.com",
            "phone": "(555) 567-8901",
            "experience": "8 years of experience",
            "skills": ["Python", "Flask", "Django", "PostgreSQL", "Redis", "Docker", "Kubernetes", "AWS", "CI/CD", "Microservices"],
            "education": "B.S. Computer Engineering, Carnegie Mellon, 2015",
            "summary": "Senior Backend Engineer with expertise in cloud infrastructure and scalable systems"
        },
        {
            "name": "Lisa Anderson",
            "filename": "lisa_anderson_resume",
            "email": "lisa.a@email.com",
            "phone": "(555) 678-9012",
            "experience": "4 years of experience",
            "skills": ["Java", "Spring", "Hibernate", "MySQL", "REST API", "JUnit", "Maven", "Jenkins"],
            "education": "B.S. Information Technology, Georgia Tech, 2019",
            "summary": "Java Developer with experience in enterprise application development"
        }
    ]
    
    os.makedirs('sample_resumes', exist_ok=True)
    
    for resume in resumes_data:
        create_docx_resume(resume)
        print(f"Created {resume['filename']}.docx")

def create_docx_resume(data):
    doc = Document()
    
    title = doc.add_heading(data['name'], 0)
    title.alignment = 1
    
    contact = doc.add_paragraph()
    contact.add_run(f"{data['email']} | {data['phone']}").italic = True
    contact.alignment = 1
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(f"{data['summary']}. {data['experience']} in the field.")
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph(data['education'])
    
    doc.add_heading('Technical Skills', level=1)
    skills_para = doc.add_paragraph()
    skills_para.add_run(', '.join(data['skills']))
    
    doc.add_heading('Experience', level=1)
    
    job1 = doc.add_paragraph()
    job1.add_run(f'Senior Developer - Tech Company (2020-Present)').bold = True
    responsibilities = [
        'Developed and maintained web applications using modern frameworks',
        'Collaborated with cross-functional teams in an Agile environment',
        'Implemented CI/CD pipelines and automated testing',
        'Mentored junior developers and conducted code reviews'
    ]
    for resp in responsibilities:
        doc.add_paragraph(resp, style='List Bullet')
    
    job2 = doc.add_paragraph()
    job2.add_run(f'Software Developer - Previous Company (2018-2020)').bold = True
    responsibilities2 = [
        'Built RESTful APIs and microservices',
        'Optimized database queries and improved application performance',
        'Participated in architecture design discussions'
    ]
    for resp in responsibilities2:
        doc.add_paragraph(resp, style='List Bullet')
    
    filename = f'sample_resumes/{data["filename"]}.docx'
    doc.save(filename)

if __name__ == '__main__':
    create_sample_resumes()
    print("\nSample resumes created successfully!")
    print("Files are located in the 'sample_resumes' folder")
